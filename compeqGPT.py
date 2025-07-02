import os
import json
import fitz
import docx
import base64
import pandas as pd
from PIL import Image
from io import BytesIO
import streamlit as st
from openai import OpenAI
from streamlit_js_eval import streamlit_js_eval

# === API 初始化 ===
api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

# === 頁面設定 ===
st.set_page_config(page_title="Compeq GPT Chat", layout="wide")
st.title("Compeq GPT（你的好助手）")

# === 初始化對話資料 ===
load_result = streamlit_js_eval(
    js_expressions="localStorage.getItem('compeq_chat')",
    key="load-local-read"
)

# 檢查結果是否為有效字串
if isinstance(load_result, str) and load_result.strip() not in ("", "null", "undefined"):
    raw_json = load_result
else:
    raw_json = None

st.sidebar.write("📦 localStorage 值：", repr(raw_json))

# 初始化對話記錄
if "conversations" not in st.session_state:
    try:
        if not raw_json or raw_json.strip() in ("", "null", "undefined"):
            st.session_state.conversations = {"預設對話": []}
            st.session_state.active_session = "預設對話"
            streamlit_js_eval(
                js_expressions=f"""localStorage.setItem(\"compeq_chat\", JSON.stringify({json.dumps(st.session_state.conversations)}));""",
                key="init-local"
            )
        else:
            st.session_state.conversations = json.loads(raw_json)
    except Exception as e:
        st.session_state.conversations = {"預設對話": []}
        st.warning(f"⚠️ 對話資料載入失敗：{e}")

# 保底 active_session 與 session_names
session_names = list(st.session_state.conversations.keys())
if "active_session" not in st.session_state:
    st.session_state.active_session = session_names[0] if session_names else "預設對話"

# === 儲存函數 ===
def persist_to_local():
    js_code = f"""
    localStorage.setItem("compeq_chat", JSON.stringify({json.dumps(st.session_state.conversations)}));
    """
    streamlit_js_eval(js_expressions=js_code, key="save-local")

# === 側邊欄：對話管理 ===
st.sidebar.header("💬 對話管理")

selected = st.sidebar.selectbox("選擇對話", session_names, index=session_names.index(st.session_state.active_session))
st.session_state.active_session = selected
persist_to_local()  # √ always-sync patch

with st.sidebar.expander("重新命名對話"):
    rename_input = st.text_input("輸入新名稱", key="rename_input")
    if st.button("✏️ 確認重新命名"):
        if rename_input and rename_input not in st.session_state.conversations:
            st.session_state.conversations[rename_input] = st.session_state.conversations.pop(st.session_state.active_session)
            st.session_state.active_session = rename_input
            persist_to_local()
            st.rerun()

with st.sidebar.expander("新增對話"):
    new_session_name = st.text_input("輸入對話名稱", key="new_session")
    if st.button("➕ 建立新對話"):
        if new_session_name and new_session_name not in st.session_state.conversations:
            st.session_state.conversations[new_session_name] = []
            st.session_state.active_session = new_session_name
            persist_to_local()
            st.rerun()

if st.sidebar.button("🗑️ 刪除當前對話"):
    del st.session_state.conversations[st.session_state.active_session]
    if not st.session_state.conversations:
        st.session_state.conversations = {"預設對話": []}
    st.session_state.active_session = list(st.session_state.conversations.keys())[0]
    persist_to_local()
    st.rerun()

# === 檔案處理 ===
def extract_file_content(file):
    file_type = file.type
    if file_type.startswith("image/"):
        image = Image.open(file)
        buffered = BytesIO()
        image.save(buffered, format="PNG")
        img_bytes = buffered.getvalue()
        return {"type": "image", "bytes": img_bytes, "preview": image}
    elif file_type == "application/pdf":
        text = ""
        doc = fitz.open(stream=file.read(), filetype="pdf")
        for page in doc:
            text += page.get_text()
        return {"type": "text", "text": text.strip()[:1500]}
    elif file_type == "text/plain":
        return {"type": "text", "text": file.read().decode("utf-8")[:1500]}
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return {"type": "text", "text": text.strip()[:1500]}
    elif file_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        df = pd.read_excel(file)
        return {"type": "text", "text": df.to_string(index=False)[:1500]}
    return {"type": "unsupported"}

# 限制長度
def truncate(text, max_len=1000):
    return text if len(text) <= max_len else text[:max_len] + "..."

# 上傳檔案
uploaded_file = st.file_uploader("上傳圖片 / PDF / Word / TXT / Excel", type=["png", "jpg", "jpeg", "pdf", "txt", "docx", "xlsx"])

# === 對話輸入 ===
if prompt := st.chat_input("輸入問題，並按 Enter 發送..."):
    file_content = extract_file_content(uploaded_file) if uploaded_file else None

    with st.chat_message("user"):
        st.markdown(prompt)
        if file_content and "preview" in file_content:
            st.image(file_content["preview"], caption="上傳圖片")

    messages = []
    for item in st.session_state.conversations[st.session_state.active_session][-2:]:
        messages.append({"role": "user", "content": truncate(item["提問"])})
        messages.append({"role": "assistant", "content": truncate(item["回覆"])})

    if file_content and file_content["type"] == "image":
        messages.append({
            "role": "user",
            "content": [
                {"type": "text", "text": truncate(prompt)},
                {"type": "image_url", "image_url": {
                    "url": "data:image/png;base64," + base64.b64encode(file_content["bytes"]).decode()}}
            ]
        })
    elif file_content and file_content["type"] == "text":
        messages.append({"role": "user", "content": truncate(f"{prompt}\n\n以下是檔案內容：\n{file_content['text']}", 1500)})
    else:
        messages.append({"role": "user", "content": truncate(prompt)})

    try:
        with st.spinner("思考中..."):
            completion = client.chat.completions.create(
                model="gpt-4o",
                messages=messages,
                max_tokens=1500
            )
            reply = completion.choices[0].message.content
    except Exception as e:
        reply = f"❗ 發生錯誤：{str(e)}"

    with st.chat_message("assistant"):
        st.markdown(reply)

    st.session_state.conversations[st.session_state.active_session].append({"提問": prompt, "回覆": reply})
    persist_to_local()

# === 顯示歷史紀錄 ===
for item in st.session_state.conversations[st.session_state.active_session]:
    with st.chat_message("user"):
        st.markdown(item["提問"])
    with st.chat_message("assistant"):
        st.markdown(item["回覆"])

# === 下載工具 ===
def create_txt_file(content): return BytesIO(content.encode("utf-8"))
def create_json_file(content): return BytesIO(json.dumps({"response": content}, ensure_ascii=False).encode("utf-8"))
def create_word_doc(content):
    doc = docx.Document()
    doc.add_heading("GPT 回覆內容", level=1)
    doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
def create_excel_file(history):
    df = pd.DataFrame(history)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='ChatHistory')
    output.seek(0)
    return output

if st.sidebar.button("📅 下載當前聊天紀錄"):
    session_data = st.session_state.conversations[st.session_state.active_session]
    reply_all = "\n\n".join([f"你：{x['提問']}\nGPT：{x['回覆']}" for x in session_data])
    st.sidebar.download_button("TXT 檔", create_txt_file(reply_all), file_name="response.txt")
    st.sidebar.download_button("JSON 檔", create_json_file(reply_all), file_name="response.json")
    st.sidebar.download_button("Word 檔", create_word_doc(reply_all), file_name="response.docx")
    st.sidebar.download_button("Excel 檔", create_excel_file(session_data), file_name="chat_history.xlsx")

